from __future__ import annotations

import argparse
import json
import math
import re
from collections import Counter
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation


SOURCE_ALIASES = {
    "date": ("日期", "预计日期", "date", "Date"),
    "store_name": ("店铺名称", "店铺", "store", "Store"),
    "country": ("国家", "country", "Country"),
    "ad_type": ("类型", "广告类型", "type", "Type"),
    "portfolio": ("广告组合", "Portfolio", "portfolio"),
    "campaign_id": ("广告活动 ID", "广告活动ID", "campaign_id", "Campaign ID"),
    "campaign_name": ("广告活动名称", "广告活动", "campaign_name", "Campaign Name", "campaign"),
    "status": ("有效状态", "广告活动状态", "status", "Status"),
    "valid": ("广告活动有效", "valid", "enabled", "Valid"),
    "is_b2b": ("isB2b", "is_b2b", "B2B"),
    "targeting_type": ("定位类型", "targeting_type", "Targeting Type"),
    "daily_budget": ("广告活动每日上限_USD", "每日预算", "daily_budget", "Daily Budget"),
    "total_budget": ("广告活动总预算_USD", "总预算", "total_budget", "Total Budget"),
    "start_date": ("广告活动开始日期", "开始日期", "start_date", "Start Date"),
    "end_date": ("广告活动结束日期", "结束日期", "end_date", "End Date"),
    "bid_strategy": ("竞价策略", "bid_strategy", "Bid Strategy"),
    "target_roas": ("目标广告支出回报率", "目标ROAS", "target_roas", "Target ROAS"),
    "store_url": ("商店网址", "店铺网址", "store_url", "Store URL"),
    "impressions": ("曝光量", "impressions", "Impressions"),
    "clicks": ("点击", "点击次数", "clicks", "Clicks"),
    "spend": ("花费-本币", "支出_USD", "Spend", "spend"),
    "sales": (
        "广告销售额-本币",
        "销售额",
        "sales",
        "Sales",
        "attributed_retail_sales_window_view_through_USD 14",
        "attributed_retail_sales_window_view_through_USD 14天",
    ),
    "orders": (
        "广告订单",
        "订单",
        "orders",
        "Orders",
        "attributed_orders_window_view_through 14",
        "attributed_orders_window_view_through 14天",
    ),
}

COMMON_BASIC_HEADERS = [
    "平台",
    "店铺",
    "投放渠道",
    "报告日期",
    "数据周期",
    "店铺名称",
    "国家",
    "类型",
    "广告组合",
    "广告活动ID",
    "isB2b",
    "定位类型",
    "广告活动名称",
    "每日预算",
    "总预算",
    "开始日期",
    "结束日期",
    "竞价策略",
    "目标ROAS",
    "店铺网址",
]

RECORD_HEADERS = ["是否调整", "调整方式", "拒绝调整原因", "调整时间"]
DEFAULT_OUTPUT_DIR = Path(r"Z:\agent_data\广告\生成报告")


class ReportGenerationError(Exception):
    pass


@dataclass(frozen=True)
class Rule:
    num: int
    condition: str
    action: str

    @property
    def text(self) -> str:
        return f"规则{self.num}: {self.condition}；动作：{self.action}"


@dataclass
class RuleResult:
    rule: Rule
    action: str
    metrics: dict[str, float | None]
    blank_reasons: dict[str, str]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate an ad adjustment report workbook.")
    parser.add_argument("--rules", required=True, help="Path to .docx or .txt rule document.")
    parser.add_argument("--data", required=True, help="Path to ad data workbook.")
    parser.add_argument("--platform", required=True)
    parser.add_argument("--store", required=True)
    parser.add_argument("--ad-channel", required=True)
    parser.add_argument("--report-date", required=True, help="YYYY-MM-DD. Resolve 'today' before calling.")
    parser.add_argument(
        "--output-dir",
        default=str(DEFAULT_OUTPUT_DIR),
        help=f"Directory for generated reports. Defaults to {DEFAULT_OUTPUT_DIR}.",
    )
    parser.add_argument(
        "--window-days",
        type=int,
        choices=(14, 30),
        default=None,
        help="Optional metric window override. By default this is inferred from the rule text.",
    )
    return parser.parse_args()


def parse_date(value: Any) -> date | None:
    if isinstance(value, datetime):
        return value.date()
    if isinstance(value, date):
        return value
    if isinstance(value, str) and value.strip():
        text = value.strip()[:10]
        for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
            try:
                return datetime.strptime(text, fmt).date()
            except ValueError:
                pass
    return None


def parse_number(value: Any) -> float:
    if value is None or value == "":
        return 0.0
    if isinstance(value, (int, float)):
        return 0.0 if isinstance(value, float) and math.isnan(value) else float(value)
    text = str(value).strip()
    if text in {"--", "有花费无销售额", "有花费无订单"}:
        return 0.0
    text = text.replace(",", "").replace("$", "").replace("￥", "")
    is_percent = text.endswith("%")
    text = text.rstrip("%").strip()
    try:
        number = float(text)
    except ValueError:
        return 0.0
    return number / 100 if is_percent else number


def read_rule_lines(path: Path) -> list[str]:
    if path.suffix.lower() == ".docx":
        document = Document(path)
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for line in cell.text.splitlines():
                        line = line.strip()
                        if line:
                            lines.append(line)
        return lines
    return [line.strip() for line in path.read_text(encoding="utf-8-sig").splitlines() if line.strip()]


def split_inline_action(text: str) -> tuple[str, str | None]:
    match = re.search(r"(?:动作|Action)\s*[:：]\s*(.+)$", text, re.IGNORECASE)
    if not match:
        return text.strip(), None
    return text[: match.start()].strip(" ；;，,"), match.group(1).strip()


def parse_rules(path: Path) -> dict[int, Rule]:
    rules: dict[int, Rule] = {}
    current_num: int | None = None
    current_condition = ""
    for line in read_rule_lines(path):
        if line in {"说明", "动作说明", "注", "备注"} or line.lower() in {"notes", "note"}:
            current_num = None
            current_condition = ""
            continue

        match = re.match(r"^(?:规则|Rule)\s*(\d+)\s*[:：.．、-]\s*(.+)$", line, re.IGNORECASE)
        if not match:
            match = re.match(r"^(\d+)\s*[.．、]\s*(.+)$", line)
        if match:
            current_num = int(match.group(1))
            current_condition, inline_action = split_inline_action(match.group(2).strip())
            if inline_action:
                rules[current_num] = Rule(current_num, current_condition, inline_action)
                current_num = None
                current_condition = ""
            continue

        action_match = re.match(r"^(?:动作|Action)\s*[:：]\s*(.+)$", line, re.IGNORECASE)
        if action_match and current_num is not None:
            rules[current_num] = Rule(current_num, current_condition, action_match.group(1).strip())
            current_num = None
            current_condition = ""

    if not rules:
        raise ReportGenerationError("RULE_PARSE_FAILED: no rules with actions were parsed from the rule document")
    return rules


def infer_window_days(rules: dict[int, Rule], override: int | None) -> int:
    if override is not None:
        return override
    text = "\n".join(f"{rule.condition} {rule.action}" for rule in rules.values()).lower()
    if "30d" in text or "30天" in text:
        return 30
    if "14d" in text or "14天" in text:
        return 14
    return 14


def infer_engine(rules: dict[int, Rule], window_days: int) -> str:
    text = "\n".join(rule.condition for rule in rules.values()).lower()
    if window_days == 30 and ("orders-30d" in text or "spend-30d" in text or "ctr" in text):
        return "orders-spend-ctr-30d"
    return "acos-14d"


def require_rules(rules: dict[int, Rule], required: range) -> None:
    missing = [num for num in required if num not in rules]
    if missing:
        raise ReportGenerationError(f"RULE_PARSE_FAILED: missing rule numbers {missing}")


def find_source_headers(headers: tuple[Any, ...]) -> dict[str, int]:
    normalized = {str(header).strip(): index for index, header in enumerate(headers) if header is not None}
    found: dict[str, int] = {}
    for key, aliases in SOURCE_ALIASES.items():
        for alias in aliases:
            if alias in normalized:
                found[key] = normalized[alias]
                break
    return found


def require_source_headers(indexes: dict[str, int], keys: list[str]) -> None:
    missing = [key for key in keys if key not in indexes]
    if missing:
        raise ReportGenerationError(f"DATA_SCHEMA_MISMATCH: missing source columns {missing}")


def load_source_rows(path: Path) -> tuple[date | None, list[dict[str, Any]], dict[str, int]]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    worksheet = workbook[workbook.sheetnames[0]]
    headers = next(worksheet.iter_rows(min_row=1, max_row=1, values_only=True))
    indexes = find_source_headers(headers)
    require_source_headers(indexes, ["date", "impressions", "clicks", "spend", "sales", "orders"])
    if "campaign_id" not in indexes and "campaign_name" not in indexes:
        raise ReportGenerationError("DATA_SCHEMA_MISMATCH: missing campaign identity column")

    latest_date: date | None = None
    rows: list[dict[str, Any]] = []
    for source_row in worksheet.iter_rows(min_row=2, values_only=True):
        parsed_date = parse_date(source_row[indexes["date"]])
        if parsed_date and (latest_date is None or parsed_date > latest_date):
            latest_date = parsed_date
        rows.append({"_source_date": parsed_date, "_raw": source_row})
    return latest_date, rows, indexes


def check_freshness(latest_date: date | None, report_date: date) -> None:
    if latest_date is None:
        raise ReportGenerationError("FRESHNESS_FAILED: no parseable source data date found")
    allowed = {report_date, report_date - timedelta(days=1)}
    if latest_date not in allowed:
        raise ReportGenerationError(
            "FRESHNESS_FAILED: "
            f"report_date={report_date.isoformat()}, "
            f"latest_source_date={latest_date.isoformat()}, "
            f"allowed={report_date.isoformat()} or {(report_date - timedelta(days=1)).isoformat()}"
        )


def value(row: tuple[Any, ...], indexes: dict[str, int], key: str, default: Any = "") -> Any:
    if key not in indexes:
        return default
    return row[indexes[key]]


def is_active(row: tuple[Any, ...], indexes: dict[str, int]) -> bool:
    active_statuses = {"enabled", "active", "有效", "启用", "投放中", "running"}
    inactive_statuses = {"paused", "暂停", "已暂停", "archived", "归档", "已归档", "disabled", "inactive"}
    if "status" in indexes:
        status = str(value(row, indexes, "status") or "").strip().lower()
        if status in inactive_statuses:
            return False
        if status in active_statuses:
            return True
        return False
    if "valid" in indexes:
        valid = str(value(row, indexes, "valid") or "").strip().lower()
        return valid in {"是", "true", "有效", "enabled", "1", "yes"}
    return True


def group_key(row: tuple[Any, ...], indexes: dict[str, int]) -> tuple[Any, ...]:
    keys = [
        "store_name",
        "country",
        "ad_type",
        "portfolio",
        "campaign_id",
        "campaign_name",
        "is_b2b",
        "targeting_type",
        "store_url",
    ]
    return tuple(value(row, indexes, key) for key in keys if key in indexes)


def new_group(
    row: tuple[Any, ...],
    indexes: dict[str, int],
    report_date: date,
    data_start: date,
    data_end: date,
    platform: str,
    store: str,
    ad_channel: str,
) -> dict[str, Any]:
    return {
        "平台": platform,
        "店铺": store,
        "投放渠道": ad_channel,
        "报告日期": report_date.isoformat(),
        "数据周期": f"{data_start.isoformat()} 至 {data_end.isoformat()}",
        "店铺名称": value(row, indexes, "store_name"),
        "国家": value(row, indexes, "country"),
        "类型": value(row, indexes, "ad_type"),
        "广告组合": value(row, indexes, "portfolio"),
        "广告活动ID": value(row, indexes, "campaign_id"),
        "isB2b": value(row, indexes, "is_b2b"),
        "定位类型": value(row, indexes, "targeting_type"),
        "广告活动名称": value(row, indexes, "campaign_name", value(row, indexes, "campaign_id")),
        "每日预算": parse_number(value(row, indexes, "daily_budget")),
        "总预算": value(row, indexes, "total_budget"),
        "开始日期": value(row, indexes, "start_date"),
        "结束日期": value(row, indexes, "end_date"),
        "竞价策略": value(row, indexes, "bid_strategy"),
        "目标ROAS": value(row, indexes, "target_roas"),
        "店铺网址": value(row, indexes, "store_url"),
        "clicks": 0.0,
        "impressions": 0.0,
        "spend": 0.0,
        "sales": 0.0,
        "orders": 0.0,
    }


def aggregate(
    rows: list[dict[str, Any]],
    indexes: dict[str, int],
    report_date: date,
    latest_date: date,
    window_days: int,
    platform: str,
    store: str,
    ad_channel: str,
) -> tuple[date, date, list[dict[str, Any]]]:
    data_end = latest_date
    data_start = data_end - timedelta(days=window_days - 1)
    groups: dict[tuple[Any, ...], dict[str, Any]] = {}
    for source in rows:
        row_date = source["_source_date"]
        row = source["_raw"]
        if row_date is None or not (data_start <= row_date <= data_end) or not is_active(row, indexes):
            continue
        key = group_key(row, indexes)
        if key not in groups:
            groups[key] = new_group(row, indexes, report_date, data_start, data_end, platform, store, ad_channel)
        group = groups[key]
        group["clicks"] += parse_number(value(row, indexes, "clicks"))
        group["impressions"] += parse_number(value(row, indexes, "impressions"))
        group["spend"] += parse_number(value(row, indexes, "spend"))
        group["sales"] += parse_number(value(row, indexes, "sales"))
        group["orders"] += parse_number(value(row, indexes, "orders"))
    return data_start, data_end, list(groups.values())


def result(rule: Rule, action: str, metrics: dict[str, float | None], blanks: dict[str, str] | None = None) -> RuleResult:
    return RuleResult(rule=rule, action=action, metrics=metrics, blank_reasons=blanks or {})


def apply_orders_spend_ctr_rules(row: dict[str, Any], rules: dict[int, Rule]) -> RuleResult | None:
    require_rules(rules, range(1, 15))
    clicks = row["clicks"]
    impressions = row["impressions"]
    spend = row["spend"]
    sales = row["sales"]
    orders = row["orders"]
    ctr = clicks / impressions if impressions > 0 else 0.0
    acos = spend / sales if sales > 0 else None
    metrics = {
        "Orders_30d": orders,
        "Spend_30d": spend,
        "Sales_30d": sales,
        "ACoS_30d": acos,
        "CTR_30d": ctr,
        "Clicks_30d": clicks,
        "Impressions_30d": impressions,
    }
    blanks = {}
    if acos is None:
        blanks["ACoS_30d空白原因"] = "近30天广告销售额为0，无法计算ACoS"

    if orders == 0:
        if spend < 120:
            rule = rules[1]
        elif spend < 220:
            rule = rules[2] if ctr < 0.005 else rules[3]
        elif spend <= 350:
            rule = rules[4] if ctr < 0.005 else rules[5]
        else:
            rule = rules[6]
        return result(rule, rule.action, metrics, blanks)

    effective_acos = acos if acos is not None else float("inf")
    if orders == 1:
        if effective_acos <= 0.22:
            rule = rules[7] if ctr >= 0.0055 else rules[8]
        else:
            rule = rules[9] if ctr >= 0.0055 else rules[10]
        return result(rule, rule.action, metrics, blanks)

    if orders >= 2:
        if effective_acos <= 0.15:
            rule = rules[11]
        elif effective_acos <= 0.25:
            rule = rules[12]
        else:
            rule = rules[13] if ctr >= 0.0055 else rules[14]
        return result(rule, rule.action, metrics, blanks)
    return None


def apply_acos_14d_rules(row: dict[str, Any], rules: dict[int, Rule]) -> RuleResult | None:
    require_rules(rules, range(1, 10))
    clicks = row["clicks"]
    impressions = row["impressions"]
    spend = row["spend"]
    sales = row["sales"]
    orders = row["orders"]
    daily_budget = row["每日预算"]
    cvr = orders / clicks if clicks > 0 else 0.0
    spend_ratio = spend / (daily_budget * 14) if daily_budget > 0 else None
    acos = spend / sales if sales > 0 else None
    effective_acos = acos
    blanks: dict[str, str] = {}
    if acos is None:
        if spend <= 0:
            return None
        blanks["ACoS_14d空白原因"] = "最近14天销售额为0，无法计算ACoS；有广告花费但无销售，按高ACoS风险判断"
        effective_acos = float("inf")
    if spend_ratio is None:
        blanks["Spend_占比_14d空白原因"] = "每日预算为空或为0，无法计算花费消耗预算占比"
    metrics = {
        "ACoS_14d": acos,
        "CVR_14d": cvr,
        "Spend_占比_14d": spend_ratio,
        "Clicks_14d": clicks,
        "Impressions_14d": impressions,
        "Spend_14d": spend,
        "Sales_14d": sales,
        "Orders_14d": orders,
    }
    assert effective_acos is not None
    if effective_acos <= 0.15:
        return result(rules[1], "增加预算" if cvr > 0.04 else "持续观察", metrics, blanks)
    if 0.15 < effective_acos <= 0.30:
        if spend_ratio is not None and spend_ratio > 1 and cvr > 0.03:
            return result(rules[2], "增加预算", metrics, blanks)
        if spend_ratio is not None and spend_ratio < 0.75 and impressions < 3000:
            return result(rules[3], "提高CPC", metrics, blanks)
        if spend_ratio is not None and spend_ratio < 0.75 and impressions > 5000:
            return result(rules[4], "点击率瓶颈，检查前台是否有划线价格", metrics, blanks)
        return None
    if 0.30 < effective_acos < 0.40:
        if clicks <= 100:
            return result(rules[5], "持续观察", metrics, blanks)
        if clicks > 100 and cvr > 0.015:
            return result(rules[6], "降低CPC", metrics, blanks)
        if clicks > 100 and cvr < 0.01:
            return result(rules[7], "暂停广告", metrics, blanks)
        return None
    if effective_acos >= 0.40:
        return result(rules[8], "暂停广告", metrics, blanks) if cvr < 0.01 else result(rules[9], "小幅降低5-10% CPC", metrics, blanks)
    return None


def apply_rules(row: dict[str, Any], rules: dict[int, Rule], engine: str) -> RuleResult | None:
    if engine == "orders-spend-ctr-30d":
        return apply_orders_spend_ctr_rules(row, rules)
    return apply_acos_14d_rules(row, rules)


def metric_headers(engine: str) -> list[str]:
    if engine == "orders-spend-ctr-30d":
        return [
            "Orders_30d",
            "Spend_30d",
            "Sales_30d",
            "ACoS_30d",
            "ACoS_30d空白原因",
            "CTR_30d",
            "Clicks_30d",
            "Impressions_30d",
        ]
    return [
        "ACoS_14d",
        "ACoS_14d空白原因",
        "CVR_14d",
        "Spend_占比_14d",
        "Spend_占比_14d空白原因",
        "Clicks_14d",
        "Impressions_14d",
        "Spend_14d",
        "Sales_14d",
        "Orders_14d",
    ]


def build_workbook(rows: list[dict[str, Any]], engine: str, report_date: date, output_path: Path) -> None:
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "广告调整报告"
    metrics = metric_headers(engine)
    headers = COMMON_BASIC_HEADERS + ["建议动作"] + ["触发规则"] + metrics + RECORD_HEADERS
    sections = [
        ("广告基本信息", COMMON_BASIC_HEADERS),
        ("根据规则调整的动作", ["建议动作"]),
        ("触发的规则", ["触发规则"]),
        ("触发规则中指标的数值", metrics),
        ("调整记录", RECORD_HEADERS),
    ]

    col = 1
    for title, section_headers in sections:
        worksheet.merge_cells(start_row=1, start_column=col, end_row=1, end_column=col + len(section_headers) - 1)
        worksheet.cell(1, col, title)
        col += len(section_headers)

    for index, header in enumerate(headers, 1):
        worksheet.cell(2, index, header)

    for row_index, row in enumerate(rows, 3):
        rule_result: RuleResult = row["result"]
        values = [row.get(header, "") for header in COMMON_BASIC_HEADERS]
        values += [rule_result.action, rule_result.rule.text]
        values += [
            rule_result.metrics.get(header, rule_result.blank_reasons.get(header, ""))
            for header in metrics
        ]
        values += ["", "", "", report_date.isoformat()]
        for col_index, cell_value in enumerate(values, 1):
            worksheet.cell(row_index, col_index, cell_value if cell_value is not None else "")

    format_workbook(worksheet, headers)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook.save(output_path)


def format_workbook(worksheet: Any, headers: list[str]) -> None:
    border = Border(
        left=Side(style="thin", color="808080"),
        right=Side(style="thin", color="808080"),
        top=Side(style="thin", color="808080"),
        bottom=Side(style="thin", color="808080"),
    )
    for row in worksheet.iter_rows(min_row=1, max_row=max(worksheet.max_row, 3), min_col=1, max_col=worksheet.max_column):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if cell.row == 1:
                cell.fill = PatternFill("solid", fgColor="D9EAF7")
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            elif cell.row == 2:
                cell.fill = PatternFill("solid", fgColor="F3F6FA")
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    worksheet.freeze_panes = "A3"
    widths = [12, 12, 15, 12, 24, 18, 10, 10, 18, 16, 8, 12, 34, 12, 12, 14, 14, 14, 12, 34]
    widths += [36, 80]
    widths += [12 for _ in range(len(headers) - len(widths) - 4)]
    widths += [12, 20, 20, 14]
    for index, width in enumerate(widths[: len(headers)], 1):
        worksheet.column_dimensions[get_column_letter(index)].width = width
    for header in {"ACoS_14d", "CVR_14d", "Spend_占比_14d", "ACoS_30d", "CTR_30d"}:
        if header in headers:
            col = headers.index(header) + 1
            for row_index in range(3, worksheet.max_row + 1):
                worksheet.cell(row_index, col).number_format = "0.00%"
    validation = DataValidation(type="list", formula1='"已调整,拒绝调整"', allow_blank=True)
    worksheet.add_data_validation(validation)
    status_col = headers.index("是否调整") + 1
    validation.add(f"{get_column_letter(status_col)}3:{get_column_letter(status_col)}{max(worksheet.max_row, 3)}")


def main() -> int:
    args = parse_args()
    report_date = datetime.strptime(args.report_date, "%Y-%m-%d").date()
    output_path = Path(args.output_dir) / f"{args.store}_{report_date.isoformat()}_广告调整报告.xlsx"
    rules = parse_rules(Path(args.rules))
    window_days = infer_window_days(rules, args.window_days)
    engine = infer_engine(rules, window_days)
    latest_date, source_rows, indexes = load_source_rows(Path(args.data))
    check_freshness(latest_date, report_date)
    assert latest_date is not None
    data_start, data_end, rows = aggregate(
        source_rows,
        indexes,
        report_date,
        latest_date,
        window_days,
        args.platform,
        args.store,
        args.ad_channel,
    )
    for row in rows:
        row["result"] = apply_rules(row, rules, engine)
    report_rows = [row for row in rows if row["result"] is not None]
    report_rows.sort(key=lambda row: (row["result"].rule.num, str(row.get("广告活动名称", ""))))
    build_workbook(report_rows, engine, report_date, output_path)
    summary = {
        "status": "success",
        "report_date": report_date.isoformat(),
        "data_period": f"{data_start.isoformat()} to {data_end.isoformat()}",
        "latest_source_date": latest_date.isoformat(),
        "window_days": window_days,
        "rule_engine": engine,
        "active_group_count": len(rows),
        "reported_row_count": len(report_rows),
        "action_counts": dict(Counter(row["result"].action for row in report_rows)),
        "output_path": str(output_path),
    }
    print(json.dumps(summary, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except ReportGenerationError as exc:
        print(json.dumps({"status": "error", "error": str(exc)}, ensure_ascii=False, indent=2))
        raise SystemExit(2)
